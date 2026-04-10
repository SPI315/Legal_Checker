from __future__ import annotations

import io

from docx import Document

from app.services.orchestration.types import Finding, PipelineResult


class DocxReportExporter:
    def export(self, result: PipelineResult) -> bytes:
        doc = Document()
        doc.add_heading("Legal Checker Report", level=1)
        doc.add_paragraph(f"Session ID: {result.session_id}")
        doc.add_paragraph(f"Status: {result.status}")
        doc.add_paragraph(f"File: {result.file_name}")
        doc.add_paragraph(f"Jurisdiction: {result.jurisdiction}")
        doc.add_paragraph(f"Findings count: {len(result.findings)}")

        if result.degraded_flags:
            doc.add_heading("Degraded Flags", level=2)
            for flag in result.degraded_flags:
                doc.add_paragraph(flag, style="List Bullet")

        doc.add_heading("Executive Summary", level=2)
        if not result.findings:
            doc.add_paragraph("No findings were produced in this run.")
        else:
            for finding in result.findings:
                doc.add_paragraph(f"{finding.title} [{finding.paragraph_id}]", style="List Bullet")

        doc.add_heading("Detailed Findings", level=2)
        for finding in result.findings:
            self._append_finding(doc, finding)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _append_finding(self, doc: Document, finding: Finding) -> None:
        doc.add_heading(f"{finding.title} ({finding.paragraph_id})", level=3)
        doc.add_paragraph(f"Summary: {finding.summary}")
        doc.add_paragraph(f"Confidence: {finding.confidence:.2f}")
        doc.add_paragraph(f"Suggested edit: {finding.suggested_edit}")
        if finding.evidence:
            doc.add_paragraph("Evidence:")
            for evidence in finding.evidence:
                doc.add_paragraph(f"Title: {evidence.title}", style="List Bullet")
                doc.add_paragraph(f"URL: {evidence.uri}")
                if evidence.snippet:
                    doc.add_paragraph(f"Snippet: {evidence.snippet}")
