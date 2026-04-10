from __future__ import annotations

import io
from pathlib import Path

from app.services.documents.types import DocumentParagraph, DocumentParseResult


class UnsupportedDocumentTypeError(ValueError):
    pass


class DocumentIngestionService:
    def parse(self, file_name: str, content: bytes) -> DocumentParseResult:
        ext = Path(file_name).suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(file_name, content)
        if ext == ".docx":
            return self._parse_docx(file_name, content)

        raise UnsupportedDocumentTypeError(
            "Unsupported file type. Allowed: .pdf, .docx"
        )

    def _parse_pdf(self, file_name: str, content: bytes) -> DocumentParseResult:
        import fitz

        doc = fitz.open(stream=content, filetype="pdf")
        paragraphs: list[DocumentParagraph] = []
        current_offset = 0

        for page_idx in range(len(doc)):
            page = doc[page_idx]
            text = page.get_text("text").strip()
            if text:
                start_offset = current_offset
                end_offset = start_offset + len(text)
                paragraphs.append(
                    DocumentParagraph(
                        paragraph_id=f"p{page_idx + 1}_1",
                        page=page_idx + 1,
                        start_offset=start_offset,
                        end_offset=end_offset,
                        text=text,
                    )
                )
                current_offset = end_offset + 2

        full_text = "\n\n".join(p.text for p in paragraphs)
        return DocumentParseResult(
            file_name=file_name,
            file_type="pdf",
            full_text=full_text,
            paragraphs=paragraphs,
        )

    def _parse_docx(self, file_name: str, content: bytes) -> DocumentParseResult:
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs: list[DocumentParagraph] = []
        idx = 0
        current_offset = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            idx += 1
            start_offset = current_offset
            end_offset = start_offset + len(text)
            paragraphs.append(
                DocumentParagraph(
                    paragraph_id=f"p{idx}",
                    page=1,
                    start_offset=start_offset,
                    end_offset=end_offset,
                    text=text,
                )
            )
            current_offset = end_offset + 2

        for table_idx, table in enumerate(doc.tables, start=1):
            for row_idx, row in enumerate(table.rows, start=1):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if not cells:
                    continue
                row_text = " | ".join(cells)
                idx += 1
                start_offset = current_offset
                end_offset = start_offset + len(row_text)
                paragraphs.append(
                    DocumentParagraph(
                        paragraph_id=f"t{table_idx}_r{row_idx}",
                        page=1,
                        start_offset=start_offset,
                        end_offset=end_offset,
                        text=row_text,
                    )
                )
                current_offset = end_offset + 2

        full_text = "\n\n".join(p.text for p in paragraphs)
        return DocumentParseResult(
            file_name=file_name,
            file_type="docx",
            full_text=full_text,
            paragraphs=paragraphs,
        )
